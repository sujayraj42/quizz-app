from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt


REPORT_PATH = r"C:\Users\sujay\OneDrive\Documents\New project\Study-Buddy_Project_Report.docx"


def set_base_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for style_name in ["Title", "Heading 1", "Heading 2"]:
        style = document.styles[style_name]
        style.font.name = "Calibri"


def add_field(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{label}: ").bold = True
    paragraph.add_run(value)


def add_bullet_list(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value


def build_report() -> None:
    document = Document()
    set_base_styles(document)

    title = document.add_paragraph()
    title.style = document.styles["Title"]
    title.alignment = 1
    title.add_run("NODE.JS PROJECT REPORT")

    subtitle = document.add_paragraph()
    subtitle.alignment = 1
    subtitle.add_run("Study-Buddy Real-Time Quiz Room").bold = True

    document.add_paragraph()
    add_field(document, "Project Title", "Study-Buddy Real-Time Quiz Room")
    add_field(document, "Student Name", "______________________________")
    add_field(document, "Roll Number", "______________________________")
    add_field(document, "Course / Semester", "______________________________")
    add_field(document, "Submitted To", "______________________________")

    document.add_page_break()

    document.add_heading("1. Problem Statement", level=1)
    document.add_paragraph(
        "Traditional classroom quizzes are often static, slow, and less engaging because "
        "students answer one by one or wait for the teacher to manually move through the flow. "
        "This reduces excitement and makes it difficult to create a synchronized quiz experience "
        "for the full class. Study-Buddy solves this problem by providing a real-time competitive "
        "quiz platform where a teacher creates a room, students join with a code, and questions "
        "are pushed simultaneously to all participants using WebSockets. The primary target users "
        "are teachers, students, and academic demonstrators who need a simple but visually strong "
        "live quiz system for presentations, class activities, or competitions."
    )

    document.add_heading("2. Objectives", level=1)
    add_bullet_list(
        document,
        [
            "Build a Node.js based real-time classroom quiz application.",
            "Allow a host to create a room and students to join instantly using a short room code.",
            "Push each question to all connected clients at the same time with Socket.io.",
            "Compute scores automatically and update the leaderboard live after every round.",
            "Provide a projector/classroom display screen for a shared presentation view.",
            "Support custom question loading from JSON so the teacher can switch quiz sets quickly.",
            "Deliver a competition-ready interface with strong visual design and mobile-friendly student access.",
        ],
    )

    document.add_heading("3. Technology Stack", level=1)
    add_bullet_list(
        document,
        [
            "Node.js: JavaScript runtime used to build the server-side application.",
            "Express.js: Web framework used for routing and serving pages/static assets.",
            "Socket.io: Real-time communication layer used to broadcast questions, timers, and leaderboard updates.",
            "EJS: Template engine used to render host, player, landing, and display pages.",
            "HTML/CSS/Vanilla JavaScript: Used for the client-side interface and interactions.",
            "In-memory JavaScript objects: Used to manage rooms, players, scores, timers, and question packs for demo simplicity.",
            "JSON files: Used for predefined quiz packs such as the India GK sample set.",
        ],
    )

    document.add_heading("4. Project Structure", level=1)
    add_table(
        document,
        ["Folder / File", "Description"],
        [
            ["server/index.js", "Main entry point that starts the Express server and registers routes."],
            ["server/socket.js", "Socket.io event handling for room creation, joining, quiz start, timer, answers, leaderboard, and display sync."],
            ["server/roomManager.js", "Core business logic for rooms, players, scoring, question packs, and quiz state management."],
            ["server/questions.js", "Default seeded question set for demo mode."],
            ["views/", "EJS pages for landing page, host screen, player screen, and classroom display screen."],
            ["public/css/styles.css", "Competition-style UI design and responsive styling."],
            ["public/js/host.js", "Host-side interactions including room control and custom question loading."],
            ["public/js/player.js", "Student-side quiz participation and answer submission logic."],
            ["public/js/display.js", "Projector/display screen logic for classroom view."],
            ["public/data/", "Sample JSON-based question packs including India GK Basics."],
        ],
    )

    document.add_heading("5. API Endpoints", level=1)
    add_table(
        document,
        ["Method", "Endpoint", "Description"],
        [
            ["GET", "/", "Landing page for project overview and player join form."],
            ["GET", "/host", "Host console used by the teacher to create and control a quiz room."],
            ["GET", "/player", "Student console for joining a room and answering questions."],
            ["GET", "/display", "Shared classroom display or projector screen for live presentation."],
            ["GET", "/health", "Simple health-check endpoint returning application status."],
        ],
    )
    document.add_paragraph(
        "Note: The most important communication in this project happens through Socket.io events "
        "instead of traditional REST APIs. Events include room:create, room:join, quiz:start, "
        "quiz:answer, quiz:next, quiz:question, quiz:timer, quiz:roundResult, quiz:final, and display:join."
    )

    document.add_heading("6. Database Models (MongoDB)", level=1)
    document.add_paragraph(
        "This project does not use MongoDB in the current version because it is optimized as a simple "
        "live demo application. Instead, it stores data in memory while the server is running. The "
        "logical data models used by the application are described below."
    )
    add_table(
        document,
        ["Logical Model", "Important Fields"],
        [
            ["Room", "code, hostSocketId, status, currentQuestionIndex, questions, players, roundAnswers, lastRound"],
            ["Player", "id, socketId, name, score, color"],
            ["Question", "id, prompt, choices, answerIndex, theme"],
            ["Question Pack", "title, source, questions[]"],
        ],
    )
    document.add_paragraph(
        "A future enhancement can replace in-memory storage with MongoDB to persist users, rooms, "
        "question banks, and historical quiz results."
    )

    document.add_heading("7. API Testing (Postman)", level=1)
    document.add_paragraph(
        "Postman testing is limited in this project because the main workflow relies on WebSocket-based "
        "real-time events rather than CRUD APIs. Functional testing was performed by running the app in "
        "multiple browser windows and validating the following scenarios:"
    )
    add_bullet_list(
        document,
        [
            "Host creates a room and receives a valid room code.",
            "Students join the room successfully using the displayed code.",
            "Host starts the quiz and all joined clients receive the same question at the same time.",
            "Timer updates are pushed live to player, host, and display screens.",
            "Answer submission is locked to one answer per player per round.",
            "Leaderboard updates instantly after each round.",
            "Custom question packs from JSON are accepted before quiz start.",
            "Projector display stays synchronized with live room progress.",
        ],
    )
    document.add_paragraph("Screenshot placeholders can be inserted here before final submission.")

    document.add_heading("8. Flow Diagram", level=1)
    document.add_paragraph(
        "System Flow:\n"
        "Teacher opens Host Console -> Creates Room -> Room Code Generated -> Students join using Player Screen -> "
        "Projector joins via Display Screen -> Teacher starts quiz -> Server broadcasts question with Socket.io -> "
        "Students submit answers -> Server evaluates correctness and score -> Leaderboard updates instantly -> "
        "Next question is broadcast -> Final podium is displayed."
    )

    document.add_page_break()

    document.add_paragraph("Suggested Flow Diagram for report:")
    document.add_paragraph(
        "[Host] -> [Node.js + Express + Socket.io Server] <- [Students]\n"
        "[Server] -> Broadcast Question -> [All Clients]\n"
        "[Students] -> Submit Answer -> [Server]\n"
        "[Server] -> Score + Leaderboard Update -> [Host / Player / Display]"
    )

    document.add_heading("9. Results", level=1)
    document.add_paragraph(
        "The Study-Buddy project was successfully implemented as a working real-time quiz platform. "
        "The host can create a room, share the room code, load a custom question pack such as the "
        "India GK Basics set, and start the quiz. Students can join from separate screens and answer "
        "questions under a live timer. The classroom display provides a shared projector-friendly view "
        "of questions and leaderboard changes. The leaderboard updates instantly without page refresh, "
        "which demonstrates the effect of WebSockets clearly during a class presentation."
    )
    add_bullet_list(
        document,
        [
            "Real-time synchronized question delivery works successfully.",
            "Leaderboard ranking changes are visible immediately after scoring.",
            "Custom question pack loading supports quick content switching.",
            "A dedicated display screen improves classroom demonstration quality.",
            "The UI is visually polished and suitable for live demo competitions.",
        ],
    )

    document.add_heading("10. Conclusion", level=1)
    document.add_paragraph(
        "Study-Buddy Real-Time Quiz Room demonstrates how Node.js and Socket.io can be used to create "
        "an engaging real-time classroom application. The project meets its main goals by enabling a host "
        "to control a synchronized quiz session, allowing multiple students to participate together, and "
        "showing instant leaderboard updates. The project is simple to run, easy to demonstrate, and "
        "effective for academic presentation."
    )
    document.add_paragraph("Future Scope:")
    add_bullet_list(
        document,
        [
            "Add persistent storage with MongoDB for question banks and result history.",
            "Introduce authentication for teachers and students.",
            "Provide an admin UI for creating and editing question packs without JSON.",
            "Add support for images, categories, and difficulty levels in questions.",
            "Generate downloadable score reports after each quiz.",
        ],
    )

    document.save(REPORT_PATH)


if __name__ == "__main__":
    build_report()
